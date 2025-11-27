import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import re

# ------------------------------------------------------------
# PRODUKTINFOS SCRAPEN
# ------------------------------------------------------------

def scrape_product_info(url):
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        modell_tag = soup.find('h1', class_='product--title') or soup.find('h1', class_='product-header-title')
        modell = modell_tag.text.strip() if modell_tag else "Modell nicht gefunden"

        artikelnummer = "Artikelnummer nicht gefunden"
        all_text = soup.get_text()
        match = re.search(r'Artikel-?Nr\.?:\s*(\d+)', all_text)
        if match:
            artikelnummer = match.group(1)

        preis_aktuell_tag = soup.find('span', class_='price--content') or soup.find('div', class_='price--current')
        preis_aktuell = preis_aktuell_tag.text.strip() if preis_aktuell_tag else None

        preis_alt_tag = soup.find('span', class_='price--line-through') or soup.find('span', class_='price-old')
        preis_alt = preis_alt_tag.text.strip() if preis_alt_tag else None

        if not preis_aktuell:
            meta_price = soup.find('meta', itemprop='price')
            if meta_price and meta_price.has_attr('content'):
                preis_aktuell = meta_price['content'].strip() + " €"

        if preis_aktuell:
            preis_aktuell = preis_aktuell.replace('*', '').strip()
        if preis_alt:
            preis_alt = preis_alt.replace('*', '').strip()

        img_url = None
        img_tag = soup.find('img', attrs={'data-img-large': True})
        if img_tag:
            img_url = img_tag['data-img-large']
        else:
            match_img = re.findall(r'data-img-large="(https://[^"]+\.jpg)"', response.text)
            if match_img:
                img_url = match_img[0]

        return modell, artikelnummer, preis_aktuell, preis_alt, img_url

    except Exception as e:
        st.error(f"Fehler beim Auslesen der Webseite: {e}")
        return None, None, None, None, None


# ------------------------------------------------------------
# WORD-DATEI ERSTELLEN (LINKS AUTOMATISCH HINTER TEXT)
# ------------------------------------------------------------

def create_word_file(modell, artikelnummer, preis_aktuell, preis_alt, img_url):
    doc = Document()
    section = doc.sections[0]

    # Seite = A4
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.orientation = WD_ORIENT.PORTRAIT

    # Normale Ränder
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(31)
    section.right_margin = Mm(31)

    # Schriftart global
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    rPr = style.element.rPr
    rFonts = rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Arial')

    # ----------------------------------------------------
    # A5-Rahmen-Tabelle
    # ----------------------------------------------------
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False

    # Innenabstände der Tabelle entfernen
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = OxmlElement("w:tblCellMar")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")
        tbl_cell_mar.append(el)
    tbl_pr.append(tbl_cell_mar)

    cell = table.rows[0].cells[0]

    # Tabelle = exakt A5
    table.columns[0].width = Mm(148)
    table.rows[0].height = Mm(210)

    # Schneide-Rahmen zeichnen
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')

    for edge in ["top", "left", "bottom", "right"]:
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:val"), "dotted")
        edge_el.set(qn("w:sz"), "10")
        edge_el.set(qn("w:color"), "C0C0C0")
        borders.append(edge_el)

    tcPr.append(borders)

    # ----------------------------------------------------
    # Hintergrundgrafik (automatisch hinter Text)
    # ----------------------------------------------------
    try:
        bg_url = "https://backend.ofen.de/media/image/63/2e/5c/Grafik-fuer-Preisschildchen-unten.png"
        bg_response = requests.get(bg_url)
        bg_response.raise_for_status()
        bg_stream = BytesIO(bg_response.content)

        # Run für Hintergrund
        p_bg = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run_bg = p_bg.add_run()
        pic = run_bg.add_picture(bg_stream, width=Mm(148), height=Mm(210))

        # XML bearbeiten, um Bild hinter Text zu legen
        pic_element = run_bg._r.xpath(".//pic:pic")[0]
        sp_pr = pic_element.xpath(".//pic:spPr")[0]
        # wp14:behindDoc hinzufügen
        behind = OxmlElement('wp14:behindDoc')
        sp_pr.insert(0, behind)

    except:
        cell.add_paragraph("Hintergrundbild konnte nicht geladen werden.")

    # ----------------------------------------------------
    # Produktbild
    # ----------------------------------------------------
    if img_url:
        try:
            img_response = requests.get(img_url)
            img_response.raise_for_status()
            img_stream = BytesIO(img_response.content)

            p_img = cell.add_paragraph()
            p_img.alignment = 1
            p_img.add_run().add_picture(img_stream, width=Mm(80))

        except:
            cell.add_paragraph("Produktbild konnte nicht geladen werden.")
    else:
        cell.add_paragraph("Kein Produktbild verfügbar.")

    # ----------------------------------------------------
    # Text
    # ----------------------------------------------------
    p = cell.add_paragraph()
    p.alignment = 1

    run1 = p.add_run(modell + "\n")
    run1.font.size = Pt(18)
    run1.font.bold = True

    run2 = p.add_run(f"Artikelnummer: {artikelnummer}\n")
    run2.font.size = Pt(11)

    run3 = p.add_run(preis_aktuell + "\n")
    run3.font.size = Pt(24)
    run3.font.bold = True
    run3.font.color.rgb = RGBColor(200, 0, 0)

    if preis_alt:
        run4 = p.add_run(preis_alt)
        run4.font.size = Pt(16)
        run4.font.strike = True
        run4.font.color.rgb = RGBColor(120, 120, 120)

    # Speichern
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# ------------------------------------------------------------
# STREAMLIT UI
# ------------------------------------------------------------

st.set_page_config(page_title="Preisschild Generator A5 mit Schneide-Rahmen", page_icon="🧾")
st.title("🧾 Preisschild Generator (A5 auf A4) mit Schneide-Rahmen")

st.markdown("**Gib den Produktlink von Ofen.de ein:**")

url = st.text_input("🔗 Produkt-URL eingeben:")

if url:
    modell, artikelnummer, preis_aktuell, preis_alt, img_url = scrape_product_info(url)

    if modell and artikelnummer and preis_aktuell:
        st.success("✅ Produktdaten erfolgreich geladen!")
        st.markdown(f"**Modell:** {modell}")
        st.markdown(f"**Artikelnummer:** {artikelnummer}")
        st.markdown(f"**Preis:** {preis_aktuell}")
        if preis_alt:
            st.markdown(f"**Alter Preis:** ~~{preis_alt}~~")
        if img_url:
            st.image(img_url, width=300)

        if st.button("📄 Preisschild erstellen"):
            file = create_word_file(modell, artikelnummer, preis_aktuell, preis_alt, img_url)
            if file:
                st.download_button(
                    label="⬇️ Preisschild als Word herunterladen",
                    data=file,
                    file_name="preisschild_A5_mit_Rahmen.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.error("❌ Einige Produktdaten konnten nicht geladen werden.")
